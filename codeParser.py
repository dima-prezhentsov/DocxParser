from docx import *
import os
from CodeFile import CodeFile


def parse(docPath):
    # doc = Document("C:\\Users\\dimon\\Desktop\\techReports\\Отчет1ЛабораторнаяПреженцовДмитрийИгоревичМ32051.docx")
    doc = Document(docPath)
    text = []

    for paragraph in doc.paragraphs:
        text.append(paragraph.text)

    # print(text)
    addCode = False
    codeFiles = []
    newCodeFile = None

    for row in range(len(text) - 1):
        splitedRow = text[row].split(' ')
        splitedNextRow = text[row + 1].split(' ')
        if addCode and newCodeFile is not None:
            if len(splitedNextRow) > 1 and splitedNextRow[0] + splitedNextRow[1] == "Добавленфайл":
                addCode = False
                codeFiles.append(newCodeFile)

            newCodeFile.addCodeRow(text[row])
        if len(splitedRow) > 1 and splitedRow[0] + splitedRow[1] == "Добавленфайл" and not addCode:
            filePath = splitedRow[2]
            dirPath = filePath.split('\\')[:-1]
            fileName = filePath.split('\\')[-1]
            addCode = True
            newCodeFile = CodeFile(dirPath, filePath)

    return codeFiles